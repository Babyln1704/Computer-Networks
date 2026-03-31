import socket
import ssl
import threading
import os
import queue
import time
from fpdf import FPDF
from PIL import Image
from docx import Document
import PyPDF2
import csv

HOST = "10.20.203.56"
PORT = 5000
BUFFER = 65536

job_queue = queue.PriorityQueue()

VALID = {
    ".txt": ["pdf", "docx", "csv"],
    ".docx": ["txt", "pdf"],
    ".pdf": ["txt", "docx"],
    ".png": ["jpg"],
    ".jpg": ["png"],
    ".jpeg": ["png"],
    ".csv": ["txt"]
}


def convert_file(filename, fmt):
    name, ext = os.path.splitext(filename)
    ext = ext.lower()
    fmt = fmt.lower()

    # TXT → PDF
    if ext == ".txt" and fmt == "pdf":
        out = name + ".pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        with open(filename, "r", encoding="utf-8") as f:
            text = f.read().encode("latin-1", "ignore").decode("latin-1")
            pdf.multi_cell(0, 10, text)

        pdf.output(out)

    # TXT → DOCX
    elif ext == ".txt" and fmt == "docx":
        out = name + ".docx"
        doc = Document()

        with open(filename, "r") as f:
            for line in f:
                doc.add_paragraph(line.strip())

        doc.save(out)

    # DOCX → TXT
    elif ext == ".docx" and fmt == "txt":
        out = name + ".txt"
        doc = Document(filename)

        with open(out, "w") as f:
            for p in doc.paragraphs:
                f.write(p.text + "\n")

    # DOCX → PDF
    elif ext == ".docx" and fmt == "pdf":
        out = name + ".pdf"
        doc = Document(filename)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for p in doc.paragraphs:
            text = p.text.encode("latin-1", "ignore").decode("latin-1")
            pdf.multi_cell(0, 10, text)

        pdf.output(out)

    # PDF → TXT
    elif ext == ".pdf" and fmt == "txt":
        out = name + ".txt"
        reader = PyPDF2.PdfReader(filename)

        with open(out, "w") as f:
            for page in reader.pages:
                f.write(page.extract_text() or "")

    # PDF → DOCX
    elif ext == ".pdf" and fmt == "docx":
        out = name + ".docx"
        reader = PyPDF2.PdfReader(filename)
        doc = Document()

        for page in reader.pages:
            doc.add_paragraph(page.extract_text() or "")

        doc.save(out)

    # IMAGE
    elif ext == ".png" and fmt == "jpg":
        out = name + ".jpg"
        Image.open(filename).save(out)

    elif ext in [".jpg", ".jpeg"] and fmt == "png":
        out = name + ".png"
        Image.open(filename).save(out)

    # CSV
    elif ext == ".csv" and fmt == "txt":
        out = name + ".txt"
        with open(filename) as f:
            reader = csv.reader(f)
            with open(out, "w") as o:
                for row in reader:
                    o.write(", ".join(row) + "\n")

    elif ext == ".txt" and fmt == "csv":
        out = name + ".csv"
        with open(filename) as f:
            lines = f.readlines()
        with open(out, "w", newline='') as o:
            writer = csv.writer(o)
            for line in lines:
                writer.writerow([line.strip()])

    else:
        return None

    return out


def worker():
    while True:
        size, filename, fmt, conn = job_queue.get()

        ext = os.path.splitext(filename)[1].lower()

        if ext not in VALID or fmt not in VALID[ext]:
            conn.send(b"ERROR")
            conn.close()
            job_queue.task_done()
            continue

        converted = convert_file(filename, fmt)

        conn.send((converted + "\n").encode())
        time.sleep(0.1)

        with open(converted, "rb") as f:
            data = f.read()

        conn.send(str(len(data)).encode().ljust(16))
        conn.sendall(data)

        conn.close()
        job_queue.task_done()


def handle_client(conn):
    header = conn.recv(1024).decode()
    filename, fmt = header.split("|")

    size = int(conn.recv(16).decode().strip())

    data = b''
    while len(data) < size:
        data += conn.recv(BUFFER)

    with open(filename, "wb") as f:
        f.write(data)

    job_queue.put((size, filename, fmt, conn))


def start_server():
    context = ssl.SSLContext(ssl.PROTOCOL_TLS_SERVER)
    context.load_cert_chain("cert.pem", "key.pem")

    server = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    server.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)

    server.bind((HOST, PORT))
    server.listen()

    print("✅ Server Running (SSL + SJF)")

    for _ in range(3):
        threading.Thread(target=worker, daemon=True).start()

    while True:
        client, _ = server.accept()
        secure = context.wrap_socket(client, server_side=True)
        threading.Thread(target=handle_client, args=(secure,)).start()


if __name__ == "__main__":
    start_server()